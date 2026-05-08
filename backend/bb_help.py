"""
BluBridge Help & Documentation module (iter67)
==============================================
Endpoints:
  GET /api/bb/help/manifest                — JSON used by the Help Center UI
  GET /api/bb/help/documentation/xlsx      — multi-sheet Excel of EVERY module + cross-cutting sections
  GET /api/bb/help/documentation/docx      — printable Word doc with TOC-style ordering
  GET /api/bb/help/template/whatsapp-resend — single-sheet upload template (kept for back-compat)
"""
import io
from fastapi import APIRouter
from fastapi.responses import StreamingResponse
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

help_router = APIRouter(prefix="/api/bb/help", tags=["Help"])

XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"
DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# ---------------------------------------------------------------------------
# Single source of truth for the Help content
# ---------------------------------------------------------------------------
MODULES = [
    {
        "id": "dashboard", "name": "Analytics Dashboard", "icon": "ChartBar", "color": "blue",
        "summary": "Upload Naukri / Pipeline / Score Sheet / College-Rank datasets and view live counts.",
        "steps": [
            "Click any of the four upload tiles to push a CSV/XLSX into the system.",
            "After upload finishes, the count badge updates automatically.",
            "Use the Candidate Journey search box to look up any applicant by email or phone.",
        ],
        "tips": [
            "Naukri must be uploaded before Pipeline (Pipeline matching depends on Naukri rows).",
            "Score Sheet is normalised so duplicates (same email or phone) are merged.",
            "Bulk upload now sends each file in its own POST sequentially — no more ingress timeouts (iter67).",
        ],
    },
    {
        "id": "hiring-forms", "name": "Hiring Forms", "icon": "FileText", "color": "violet",
        "summary": "Create custom application forms with conditional logic for each job role.",
        "steps": [
            "Click '+ New Form' and pick the form type.",
            "Add fields, options, and conditional rules.",
            "Toggle 'Show Instruction Page' if you want a pre-form instruction screen.",
            "Copy the public URL and share with candidates.",
        ],
        "tips": [
            "Slug-based URLs: /register/<your-slug>",
            "On submit, hr_team is auto-populated from form_type_name (iter67 fix #3).",
        ],
    },
    {
        "id": "interview-reports", "name": "Interview Schedule Reports", "icon": "CalendarCheck", "color": "cyan",
        "summary": "Filterable report of every scheduled interview with attendance + college-tier breakdown.",
        "steps": [
            "Filter by date range, job role, attendance, or college type.",
            "Click 'Export' to pick fields to include in the CSV.",
        ],
        "tips": ["The KPI strip (Attended / Not Attended / Premium / Non-Premium) reflects current filters."],
    },
    {
        "id": "update-scores", "name": "Update Applicants Scores", "icon": "PencilLine", "color": "emerald",
        "summary": "Manage interview rounds and update each candidate's per-round score, command, status.",
        "steps": [
            "Pick a candidate row and click 'Update Score'.",
            "Add or edit round entries (Round name, Score, Command, Status).",
            "Save — changes are append-only; previous values are preserved in history.",
        ],
        "tips": ["Use 'Manage Rounds' to add or rename round columns globally."],
    },
    {
        "id": "score-round", "name": "Score & Round", "icon": "Table", "color": "sky",
        "summary": "Excel-like grid of every candidate × round score + induction dates.",
        "steps": [
            "Use the top filter bar to narrow by Job Role / Status / Date range.",
            "Click any cell to inline-edit. Status filter shows only Shortlisted, Rejected, On-Hold (typos auto-grouped).",
        ],
        "tips": [
            "Three induction date columns — DOJ, Onboarding, Termination — are dedicated pickers.",
            "Export now includes ALL rounds even when a date filter is applied (iter67 fix #6).",
        ],
    },
    {
        "id": "candidate-journey", "name": "Candidate Journey", "icon": "MagnifyingGlass", "color": "indigo",
        "summary": "End-to-end timeline for a single candidate (rounds, scores, status, induction).",
        "steps": ["Search by email or phone.", "View the timeline + every recorded score change."],
        "tips": ["Useful for HR queries — 'why was this candidate rejected?'"],
    },
    {
        "id": "whatsapp-resend", "name": "WhatsApp Missed Schedule Link Resend", "icon": "WhatsappLogo", "color": "whatsapp",
        "summary": "Re-deliver interview schedules + meeting links via WhatsApp for candidates who missed/deleted their original message.",
        "steps": [
            "Download the .xlsx template (Help page or upload zone).",
            "Fill Name/Email/Phone, save, and upload.",
            "Auto-match against pipeline_data + bb_registrations runs instantly.",
            "Preview the table — sendable rows have an active schedule + valid phone.",
            "Click 'Bulk Resend' or per-row send.",
            "Track every send in 'Resend History'.",
        ],
        "tips": [
            "5-min cooldown per candidate.",
            "Strict allowlist applies — non-allowlisted recipients log as 'blocked'.",
            "Reuses the AiSensy 'Candidate FollowUp' template (5 params).",
            "Excel/CSV export of the preview table is available post-upload (iter67).",
        ],
        "downloads": [{"label": "WhatsApp Resend Template (.xlsx)", "url": "/api/bb/help/template/whatsapp-resend"}],
    },
    {
        "id": "manual-alerts", "name": "Manual Applicant Alerts", "icon": "EnvelopeSimple", "color": "violet",
        "summary": "Manually re-fire any of 5 messaging flows (shortlist / schedule detail / OTP / follow-up / reject) for a single applicant.",
        "steps": [
            "Enter Email + Phone of the candidate.",
            "Click 'Search' — applicant details load below.",
            "Click any of the 5 colored action buttons to trigger that template (Mail + WhatsApp).",
        ],
        "tips": [
            "Outbound is gated by the messaging allowlist — non-allowlisted recipients are blocked at the send layer.",
            "If no schedule_token exists, one is auto-generated for the shortlist/follow-up flows.",
        ],
    },
    {
        "id": "manual-otp-verify", "name": "Manual OTP Verify", "icon": "ShieldCheck", "color": "orange",
        "summary": "Mark a candidate as Attended (otp_verified=true) by matching email + phone.",
        "steps": [
            "Enter both Email and Phone (must belong to the SAME pipeline_data row).",
            "Click 'Verify' — record is updated and a result card shows the applicant details.",
        ],
        "tips": ["Mirrors the update to bb_registrations so dashboards stay aligned."],
    },
    {
        "id": "tester-credentials", "name": "Tester Credentials", "icon": "Flask", "color": "pink",
        "summary": "Manage QA email/phone pairs that bypass the 4-month registration cooldown.",
        "steps": [
            "Add a tester (Email + Phone).",
            "Edit / Delete via the cards grid.",
            "When a candidate registers, if their email OR phone matches a tester, the cooldown is skipped.",
        ],
        "tips": [
            "Default seed: rishi.nayak@blubridge.com / 9443109903 and rajlearn@gmail.com / 8883847098.",
            "Duplicates are rejected by email-OR-phone match.",
            "This collection is ADDITIONAL to the hard-coded allowlist in messaging.py — that allowlist still gates outbound messages.",
        ],
    },
    {
        "id": "manage-job-roles", "name": "Create Job Roles", "icon": "Briefcase", "color": "navy",
        "summary": "Define and manage job titles used everywhere in the app.",
        "steps": ["Click 'Add Role', enter the title, click 'Save'."],
        "tips": ["Renaming a role updates references in pipeline_data automatically."],
    },
    {
        "id": "job-openings", "name": "Create Job Openings", "icon": "FolderOpen", "color": "rose",
        "summary": "Publish job openings with description, requirements, and status.",
        "steps": ["Add opening → fill description/requirements → toggle 'Active' to publish."],
        "tips": [],
    },
    {
        "id": "college-schedules", "name": "College Drives", "icon": "GraduationCap", "color": "pink",
        "summary": "Configure interview schedules per college and role.",
        "steps": ["Add a college, pick the job role chips, set the date/time, save."],
        "tips": ["Each college schedule generates a public registration link."],
    },
    {
        "id": "set-holidays", "name": "Set Holidays", "icon": "CalendarBlank", "color": "orange",
        "summary": "Block specific dates from being chosen as interview slots.",
        "steps": ["Pick a date → enter a name → save. Public schedule pages will hide that date."],
        "tips": [],
    },
    {
        "id": "verify-otp", "name": "Verify Applicant OTP", "icon": "ShieldCheck", "color": "teal",
        "summary": "Confirm applicant attendance via phone-number OTP at reception.",
        "steps": [
            "Enter the candidate's phone (last 10 digits) and OTP.",
            "On success, the candidate's status updates to 'Attended' across pipeline_data + bb_registrations.",
        ],
        "tips": ["OTP window: 3 hours before the scheduled time → 8 hours after it expires."],
    },
    {
        "id": "help", "name": "Help Center", "icon": "Question", "color": "navy",
        "summary": "This documentation hub — guides + FAQs + downloadable Excel/Word documentation.",
        "steps": [
            "Use the sidebar to jump to any module's guide.",
            "Click 'Download Documentation (.xlsx)' for the full multi-sheet Excel.",
            "Click 'Download Documentation (.docx)' for a printable Word version.",
        ],
        "tips": ["Documentation is generated server-side from the same manifest the UI consumes — always in sync."],
    },
]

WHATS_NEW = [
    ("Cream Light Theme",     "Full app converted from dark mode to a warm cream theme with persistent left sidebar."),
    ("WhatsApp Resend Module","CSV/XLSX upload → 5-priority match → preview → bulk send via approved 'Candidate FollowUp' template."),
    ("Manual Applicant Alerts","Manually fire any of 5 messaging flows for a single candidate (shortlist / schedule / OTP / follow-up / reject)."),
    ("Manual OTP Verify",     "Email + Phone match → set otp_verified=true with full applicant readout."),
    ("Tester Credentials",    "QA email/phone pairs bypass the 4-month cooldown without touching code."),
    ("Bulk Upload — Sequential","Each file uploads in its own POST. K8s ingress timeouts eliminated."),
    ("Bulk Upload — Cross-host fix","Per-host claim isolation (HOST_ID + queued_local status) prevents another deployment's worker from sniping our jobs."),
    ("Dedup — email OR phone", "Registration finds existing by email OR phone, updates the SAME _id (no duplicates)."),
    ("hr_team auto-populated", "Set from form_type_name on every submission."),
    ("Export — all round columns","bb_rounds is the source of truth — date filter no longer drops columns."),
    ("Help Center",           "Multi-sheet Excel + Word documentation, downloadable from the sidebar."),
]

MATCH_PRIORITY = [
    ("P1", "Exact Name + Email", "100%", "Exact Match"),
    ("P2", "Exact Name + Phone (last 10 digits)", "95%",  "Exact Match"),
    ("P3", "Email only — single candidate found", "80%",  "Partial Match"),
    ("P3", "Email only — multiple candidates found", "75%",  "Multiple Match"),
    ("P4", "Phone only — single candidate found", "75%",  "Partial Match"),
    ("P4", "Phone only — multiple candidates found", "70%",  "Multiple Match"),
    ("—",  "No match in pipeline_data or bb_registrations", "0%", "No Match"),
]

ALLOWLIST_NOTES = [
    "Strict allowlist defined in /app/backend/messaging.py (_ALLOWED_PAIRS).",
    "Default allowlist: (rishi.nayak@blubridge.com, 9443109903) and (rajlearn@gmail.com, 8883847098).",
    "All outbound WhatsApp + Email pass through is_allowed_recipient() — non-allowlisted pairs are silently dropped (logged).",
    "Tester Credentials (bb_test_credentials) bypass the 4-month REGISTRATION cooldown only — they DO NOT bypass the messaging allowlist.",
]

GLOBAL_FAQS = [
    ("Why does WhatsApp say 'blocked'?",
     "Strict allowlist enforces only configured (email, phone) pairs receive real outbound messages."),
    ("Why does my upload show 0 matched?",
     "Column headers may not match expected aliases. Use the .xlsx template — rename headers to Name / Email / Phone."),
    ("How do I add a new accepted recipient?",
     "Edit /app/backend/messaging.py → _ALLOWED_PAIRS tuple."),
    ("Can I export a table?",
     "Most module tables support 'Export' with field selection. WhatsApp Resend offers Excel + CSV export of the preview."),
    ("What is the difference between Roles and Attended Roles?",
     "Roles = all candidates. Attended Roles = candidates whose otp_verified=true (showed up at reception)."),
    ("Why do I see duplicate applicants?",
     "Pre-iter67 dedup matched only on email AND phone. iter67 fix uses email OR phone — re-register a duplicate to consolidate; old duplicates need manual cleanup."),
    ("hr_team is null on old records — why?",
     "Auto-population was added in iter67. New submissions populate it; back-fill old rows by re-submitting via the same hiring form."),
    ("Bulk upload says 'File not found on disk' — what gives?",
     "Pre-iter67 the queue was shared with another deployment that had a different filesystem. Now per-host isolation prevents this. Click 'Clear All' on Failed if any legacy rows remain."),
    ("Can I bypass cooldown for testing?",
     "Yes — add the tester's email/phone via Tester Credentials. ANY match (email or phone) bypasses the 4-month restriction."),
    ("Where do the 5 manual alert templates come from?",
     "ShortList / Schedule Detail / OTP With Job / Candidate FollowUp / Reject — all pre-approved on AiSensy."),
    ("Why does the OTP page have a 3h window?",
     "OTP becomes valid 3h before the scheduled time and stays valid for 8h after."),
    ("How are status typos handled in Score & Round?",
     "'Shortlsited', 'rejeceted', 'Hold', 'On hold' are auto-grouped into Shortlisted / Rejected / On-Hold for filtering."),
    ("Where is the WhatsApp template defined?",
     "The text body shown in the Help / preview is informational. Real WhatsApp messages go through AiSensy's pre-approved campaign templates."),
    ("Can I download the entire documentation?",
     "Yes — Help Center top-right offers .xlsx and .docx exports of every module + cross-cutting topics."),
    ("How do I roll back?",
     "Use Emergent's 'Rollback' option in the chat input — code-level rollback only, doesn't touch DB."),
]

GLOSSARY = [
    ("pipeline_data",       "Master HR table. Fields: name, email, phone, job_role, schedule_date, schedule_time, status, otp, otp_verified, hr_team, scores[], result_status, ..."),
    ("bb_registrations",    "Public registrations + reschedule tokens. schedule_token feeds the public /schedule-interview/<token> URL."),
    ("bb_rounds",           "List of interview rounds (active=true). Source of truth for export round columns."),
    ("bb_form_types",       "Form types — referenced by bb_hiring_forms and used to populate pipeline_data.hr_team."),
    ("bb_hiring_forms",     "Public form definitions with conditional logic and slug-based URLs."),
    ("bb_test_credentials", "(iter67) QA email/phone pairs that bypass the 4-month registration cooldown."),
    ("bb_resend_uploads",   "(iter67) WhatsApp Resend upload bundles — parsed rows + match data."),
    ("bb_resend_history",   "(iter67) Per-send log for WhatsApp Resend (status, retries, failure reason)."),
    ("bulk_upload_queue",   "Bulk upload queue. Statuses: queued_local (private), processing, completed, failed, archived."),
    ("HOST_ID",             "(iter67) hostname-derived stamp on each upload row preventing cross-host worker conflicts."),
]

API_REFERENCE = [
    ("GET",  "/api/bb/score-round/table",          "Excel-like grid (filters: status, date range, job role)"),
    ("GET",  "/api/bb/score-round/result-statuses","Distinct status buckets (3 main groups)"),
    ("GET",  "/api/bb/score-round/export",         "Download grid CSV (all round columns)"),
    ("POST", "/api/bb/resend/upload",              "WhatsApp Resend — parse + auto-match"),
    ("GET",  "/api/bb/resend/preview/{id}",        "Paginated preview of an upload"),
    ("POST", "/api/bb/resend/send",                "Bulk / single resend"),
    ("GET",  "/api/bb/resend/export/{id}",         "Excel/CSV export of preview rows"),
    ("GET",  "/api/bb/resend/history",             "Resend history log"),
    ("GET",  "/api/bb/manual/applicant/lookup",    "Manual Alerts — lookup applicant by email/phone"),
    ("POST", "/api/bb/manual/alerts/send-shortlist","Fire shortlist mail+WhatsApp"),
    ("POST", "/api/bb/manual/alerts/send-schedule-detail","Fire schedule detail mail+WhatsApp"),
    ("POST", "/api/bb/manual/alerts/send-otp",     "Fire OTP mail+WhatsApp"),
    ("POST", "/api/bb/manual/alerts/send-followup","Fire follow-up mail+WhatsApp"),
    ("POST", "/api/bb/manual/alerts/send-reject",  "Fire rejection mail+WhatsApp"),
    ("POST", "/api/bb/manual/otp/verify",          "Manual OTP verify (email + phone match)"),
    ("GET",  "/api/bb/manual/test-credentials",    "List tester pairs"),
    ("POST", "/api/bb/manual/test-credentials",    "Add tester pair"),
    ("PUT",  "/api/bb/manual/test-credentials/{id}","Update tester pair"),
    ("DELETE","/api/bb/manual/test-credentials/{id}","Delete tester pair"),
    ("GET",  "/api/bb/help/manifest",              "JSON manifest used by the Help Center UI"),
    ("GET",  "/api/bb/help/documentation/xlsx",    "Multi-sheet Excel of full documentation"),
    ("GET",  "/api/bb/help/documentation/docx",    "Word document of full documentation"),
    ("POST", "/api/bulk-upload/{type}",            "Bulk upload (one file per call — sequential)"),
    ("GET",  "/api/bulk-upload/status",            "Per-type pending/processed/failed lists"),
    ("POST", "/api/bulk-upload/{type}/clear-failed","Archive all failed rows for a type"),
]

# ============================================================================
# XLSX builder
# ============================================================================
_HDR_FILL    = PatternFill(start_color="1D3A8A", end_color="1D3A8A", fill_type="solid")
_HDR_FONT    = Font(bold=True, color="FFFFFF", size=11)
_SECT_FILL   = PatternFill(start_color="FAF9F1", end_color="FAF9F1", fill_type="solid")
_SECT_FONT   = Font(bold=True, color="1A2332", size=12)
_BORDER      = Border(*[Side(style="thin", color="E5E3D8")] * 4)


def _hdr(ws, row, headers):
    for i, h in enumerate(headers, 1):
        c = ws.cell(row=row, column=i, value=h)
        c.fill = _HDR_FILL; c.font = _HDR_FONT
        c.alignment = Alignment(horizontal="left", vertical="center", wrap_text=True)
        c.border = _BORDER


def _section(ws, row, title, span=4):
    c = ws.cell(row=row, column=1, value=title)
    c.fill = _SECT_FILL; c.font = _SECT_FONT
    ws.merge_cells(start_row=row, start_column=1, end_row=row, end_column=span)


def _autosize(ws, widths):
    for i, w in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = w


def _build_xlsx() -> bytes:
    wb = Workbook()

    # ---- Sheet 1: Overview / Index ----
    ws = wb.active
    ws.title = "Overview"
    _section(ws, 1, "BluBridge Hiring Pipeline — Documentation", span=4)
    _hdr(ws, 3, ["#", "Module", "Summary", "Sheet"])
    for i, m in enumerate(MODULES, 4):
        ws.cell(row=i, column=1, value=i - 3).border = _BORDER
        ws.cell(row=i, column=2, value=m["name"]).border = _BORDER
        ws.cell(row=i, column=3, value=m["summary"]).border = _BORDER
        ws.cell(row=i, column=4, value="Modules").border = _BORDER
    _autosize(ws, [6, 38, 70, 14])

    # ---- Sheet 2: Modules ----
    ws2 = wb.create_sheet("Modules")
    _section(ws2, 1, "Per-Module Guides", span=4)
    _hdr(ws2, 3, ["Module", "Summary", "How to use", "Tips"])
    for i, m in enumerate(MODULES, 4):
        ws2.cell(row=i, column=1, value=m["name"]).border = _BORDER
        ws2.cell(row=i, column=2, value=m["summary"]).border = _BORDER
        steps = "\n".join(f"{n}. {s}" for n, s in enumerate(m.get("steps", []), 1))
        ws2.cell(row=i, column=3, value=steps).border = _BORDER
        tips = "\n".join(f"• {t}" for t in m.get("tips", []))
        ws2.cell(row=i, column=4, value=tips).border = _BORDER
        for col in range(1, 5):
            ws2.cell(row=i, column=col).alignment = Alignment(wrap_text=True, vertical="top")
        ws2.row_dimensions[i].height = max(60, 18 * (1 + len(m.get("steps", [])) + len(m.get("tips", []))))
    _autosize(ws2, [30, 50, 65, 50])
    ws2.freeze_panes = "A4"

    # ---- Sheet 3: What's New ----
    ws3 = wb.create_sheet("What's New")
    _section(ws3, 1, "Latest Changes (iter67)", span=2)
    _hdr(ws3, 3, ["Item", "Description"])
    for i, (item, desc) in enumerate(WHATS_NEW, 4):
        ws3.cell(row=i, column=1, value=item).border = _BORDER
        ws3.cell(row=i, column=2, value=desc).border = _BORDER
        ws3.cell(row=i, column=2).alignment = Alignment(wrap_text=True)
    _autosize(ws3, [32, 90])

    # ---- Sheet 4: Match Logic ----
    ws4 = wb.create_sheet("Match Logic")
    _section(ws4, 1, "WhatsApp Resend — 5-Priority Match", span=4)
    _hdr(ws4, 3, ["Priority", "Rule", "Confidence", "Status"])
    for i, row in enumerate(MATCH_PRIORITY, 4):
        for col, val in enumerate(row, 1):
            ws4.cell(row=i, column=col, value=val).border = _BORDER
    _autosize(ws4, [10, 60, 14, 18])

    # ---- Sheet 5: Allowlist & Cooldown ----
    ws5 = wb.create_sheet("Allowlist & Cooldown")
    _section(ws5, 1, "Outbound Allowlist & Cooldown Bypass", span=2)
    _hdr(ws5, 3, ["#", "Rule"])
    for i, n in enumerate(ALLOWLIST_NOTES, 4):
        ws5.cell(row=i, column=1, value=i - 3).border = _BORDER
        ws5.cell(row=i, column=2, value=n).border = _BORDER
        ws5.cell(row=i, column=2).alignment = Alignment(wrap_text=True)
    _autosize(ws5, [6, 110])

    # ---- Sheet 6: FAQs ----
    ws6 = wb.create_sheet("FAQs")
    _section(ws6, 1, "Frequently Asked Questions", span=2)
    for i, (q, a) in enumerate(GLOBAL_FAQS, 0):
        rq = 3 + i * 2
        ws6.cell(row=rq, column=1, value=f"Q{i+1}: {q}").font = Font(bold=True, color="1A2332")
        ws6.cell(row=rq + 1, column=1, value=f"A: {a}").alignment = Alignment(wrap_text=True, vertical="top")
        ws6.row_dimensions[rq + 1].height = 45
    _autosize(ws6, [120])

    # ---- Sheet 7: Glossary ----
    ws7 = wb.create_sheet("Glossary")
    _section(ws7, 1, "Glossary — Mongo Collections & Concepts", span=2)
    _hdr(ws7, 3, ["Term", "Description"])
    for i, (t, d) in enumerate(GLOSSARY, 4):
        ws7.cell(row=i, column=1, value=t).border = _BORDER
        ws7.cell(row=i, column=2, value=d).border = _BORDER
        ws7.cell(row=i, column=2).alignment = Alignment(wrap_text=True)
    _autosize(ws7, [22, 110])

    # ---- Sheet 8: API Reference ----
    ws8 = wb.create_sheet("API Reference")
    _section(ws8, 1, "API Endpoints (selected)", span=3)
    _hdr(ws8, 3, ["Method", "Path", "Purpose"])
    for i, (mtd, path, purpose) in enumerate(API_REFERENCE, 4):
        ws8.cell(row=i, column=1, value=mtd).border = _BORDER
        ws8.cell(row=i, column=2, value=path).border = _BORDER
        ws8.cell(row=i, column=3, value=purpose).border = _BORDER
    _autosize(ws8, [10, 50, 70])

    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return buf.read()


# ============================================================================
# DOCX builder
# ============================================================================
def _add_h(doc, text, level=1, color="1A2332"):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor.from_string(color)
        run.font.name = "Calibri"


def _add_p(doc, text, bold=False, italic=False, color=None, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def _build_docx() -> bytes:
    doc = Document()

    # Cover
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tr = title.add_run("BluBridge Hiring Pipeline")
    tr.bold = True; tr.font.size = Pt(28); tr.font.color.rgb = RGBColor.from_string("1D3A8A")
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Complete Documentation · iter67")
    sr.font.size = Pt(14); sr.font.color.rgb = RGBColor.from_string("6B7280")
    doc.add_paragraph()

    # Section: What's New
    _add_h(doc, "What's New (iter67)", level=1)
    for item, desc in WHATS_NEW:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{item} — "); r.bold = True
        p.add_run(desc)

    doc.add_page_break()

    # Section: Modules
    _add_h(doc, "Modules", level=1)
    for m in MODULES:
        _add_h(doc, m["name"], level=2, color="1D3A8A")
        _add_p(doc, m["summary"], italic=True, color="6B7280")
        if m.get("steps"):
            _add_p(doc, "How to use", bold=True)
            for i, s in enumerate(m["steps"], 1):
                doc.add_paragraph(f"{i}. {s}", style="List Number")
        if m.get("tips"):
            _add_p(doc, "Tips", bold=True)
            for t in m["tips"]:
                doc.add_paragraph(t, style="List Bullet")
        if m.get("downloads"):
            _add_p(doc, "Downloads", bold=True)
            for d in m["downloads"]:
                p = doc.add_paragraph(style="List Bullet")
                p.add_run(f"{d['label']}  →  ").bold = True
                p.add_run(d["url"])
        doc.add_paragraph()

    doc.add_page_break()

    # Section: Match Priority Logic
    _add_h(doc, "WhatsApp Resend — 5-Priority Match Logic", level=1)
    table = doc.add_table(rows=1, cols=4)
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(["Priority", "Rule", "Confidence", "Status"]):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
    for row in MATCH_PRIORITY:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)

    doc.add_page_break()

    # Section: Allowlist & Cooldown
    _add_h(doc, "Outbound Allowlist & Cooldown Bypass", level=1)
    for n in ALLOWLIST_NOTES:
        doc.add_paragraph(n, style="List Bullet")

    doc.add_page_break()

    # Section: Glossary
    _add_h(doc, "Glossary", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Term"; hdr_cells[1].text = "Description"
    for run in hdr_cells[0].paragraphs[0].runs: run.bold = True
    for run in hdr_cells[1].paragraphs[0].runs: run.bold = True
    for term, desc in GLOSSARY:
        cells = table.add_row().cells
        cells[0].text = term; cells[1].text = desc

    doc.add_page_break()

    # Section: API Reference
    _add_h(doc, "API Reference (selected)", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Light Grid Accent 1"
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(["Method", "Path", "Purpose"]):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs: run.bold = True
    for mtd, path, purpose in API_REFERENCE:
        cells = table.add_row().cells
        cells[0].text = mtd; cells[1].text = path; cells[2].text = purpose

    doc.add_page_break()

    # Section: FAQs
    _add_h(doc, "Frequently Asked Questions", level=1)
    for i, (q, a) in enumerate(GLOBAL_FAQS, 1):
        _add_p(doc, f"Q{i}. {q}", bold=True, color="1D3A8A")
        _add_p(doc, f"A. {a}")
        doc.add_paragraph()

    buf = io.BytesIO(); doc.save(buf); buf.seek(0)
    return buf.read()


# ============================================================================
# Endpoints
# ============================================================================
@help_router.get("/manifest")
async def help_manifest():
    """Manifest the Help Center UI consumes."""
    return {
        "modules": MODULES,
        "global_faqs": [{"q": q, "a": a} for q, a in GLOBAL_FAQS],
        "whats_new": [{"item": i, "desc": d} for i, d in WHATS_NEW],
        "downloads": [
            {"label": "Full Documentation (.xlsx)", "url": "/api/bb/help/documentation/xlsx", "format": "xlsx"},
            {"label": "Full Documentation (.docx)", "url": "/api/bb/help/documentation/docx", "format": "docx"},
        ],
    }


@help_router.get("/documentation/xlsx")
async def download_documentation_xlsx():
    data = _build_xlsx()
    return StreamingResponse(
        io.BytesIO(data), media_type=XLSX_MIME,
        headers={"Content-Disposition": 'attachment; filename="BluBridge-Documentation.xlsx"'},
    )


@help_router.get("/documentation/docx")
async def download_documentation_docx():
    data = _build_docx()
    return StreamingResponse(
        io.BytesIO(data), media_type=DOCX_MIME,
        headers={"Content-Disposition": 'attachment; filename="BluBridge-Documentation.docx"'},
    )


@help_router.get("/template/whatsapp-resend")
async def download_whatsapp_resend_template():
    """Single-sheet Excel template for the WhatsApp Resend upload — kept for back-compat."""
    wb = Workbook()
    ws = wb.active; ws.title = "Template"
    _hdr(ws, 1, ["Name", "Email", "Phone"])
    for r_idx, row in enumerate([
        ["Rishi Nayak", "rishi.nayak@blubridge.com", "9443109903"],
        ["Sharmila R",  "sharmilaramu772@gmail.com", "8015527422"],
        ["Aravind K",   "aravind.k@example.com",     "9876543210"],
    ], 2):
        for c_idx, val in enumerate(row, 1):
            ws.cell(row=r_idx, column=c_idx, value=val).border = _BORDER
    ws.cell(row=6, column=1, value="Notes:").font = Font(bold=True, color="1A2332")
    notes = [
        "1. Replace example rows with your candidates.",
        "2. Column names auto-mapped: name | candidate_name | full_name · email | email_id | mail · phone | mobile | contact",
        "3. Phone must be a 10-digit Indian number starting with 6/7/8/9.",
    ]
    for i, n in enumerate(notes):
        ws.cell(row=7 + i, column=1, value=n).font = Font(color="3F4655")
    _autosize(ws, [22, 32, 18])
    buf = io.BytesIO(); wb.save(buf); buf.seek(0)
    return StreamingResponse(
        buf, media_type=XLSX_MIME,
        headers={"Content-Disposition": 'attachment; filename="BluBridge-WhatsApp-Resend-Template.xlsx"'},
    )
